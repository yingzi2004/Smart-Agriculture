from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "doc" / "智慧农业项目开发说明.docx"


BLUE = RGBColor(31, 77, 120)
GREEN = RGBColor(35, 110, 72)
MUTED = RGBColor(90, 96, 106)


def set_east_asia(run, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_run(paragraph, text, size=11, bold=False, color=None):
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, value in {"top": "90", "bottom": "90", "start": "120", "end": "120"}.items():
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, GREEN),
        ("Heading 3", 12, BLUE),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "智慧农业项目开发说明", size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    add_run(p, "成员阅读版：说明系统做什么、各模块怎么做、五个人怎么分工", size=11, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, "说明：", bold=True, color=GREEN)
    add_run(p, "这份文档不是完整论文，也不是详细需求规格书，重点是让项目成员快速统一理解，后续按模块开发。")


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text):
    p = doc.add_paragraph()
    add_run(p, text)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, item, size=10.5)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(hdr[i])
        set_cell_shading(hdr[i], "E8F1EC")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, size=10, bold=True, color=BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, text, size=9.5)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def build():
    doc = Document()
    configure_doc(doc)
    title(doc)

    h1(doc, "第1章 项目概述")
    para(doc, "本项目是一个智慧农业管理系统，分为管理员端和农户端。管理员端负责基地、大棚、作物、设备、告警、任务、数据和社区的统一管理；农户端负责查看自己负责的大棚、处理任务、接收告警、上传记录和使用 AI 建议。")
    h2(doc, "1.1 背景与目标")
    bullets(doc, [
        "解决传统农业数据分散、人工巡检效率低、异常发现不及时的问题。",
        "通过传感器、后台管理、AI 分析和农户端任务闭环，把种植过程数字化。",
        "先完成可演示 Demo，再逐步接入真实设备、真实模型和真实业务数据。",
    ])
    h2(doc, "1.2 技术选型")
    table(doc, ["类型", "技术"], [
        ["后端", "Spring Boot、Shiro、MyBatis、MySQL"],
        ["前端", "RuoYi 模板、Bootstrap、jQuery、ECharts"],
        ["AI", "LangChain、Agent、Tool、Memory、RAG 知识库"],
        ["部署", "本地开发先用单体服务，后续可拆分 AI 服务和业务服务"],
    ], [1.2, 5.4])
    h2(doc, "1.3 创新点")
    bullets(doc, [
        "管理员端和农户端分角色展示，避免农户看到复杂后台功能。",
        "环境异常、告警、农事任务形成闭环：发现问题、派发任务、执行反馈、关闭告警。",
        "引入 LangChain 智能中枢，让 AI 能读取数据、调用工具并给出种植建议。",
        "病虫害识别、RAG 知识库和社区问答结合，沉淀项目知识。",
    ])

    h1(doc, "第2章 系统总体架构")
    para(doc, "系统可以理解为四层：用户页面层、业务服务层、AI 智能中枢层、数据存储层。")
    table(doc, ["层级", "作用", "包含内容"], [
        ["用户页面层", "给不同角色使用", "管理员端、农户端、公开溯源页"],
        ["业务服务层", "处理系统核心业务", "种植、环境、告警、任务、溯源、社区"],
        ["AI 智能中枢层", "负责智能分析和工具调用", "Agent、Tool、Memory、RAG、Prompt、Workflow"],
        ["数据存储层", "保存业务数据和知识数据", "MySQL、传感器数据、向量库、日志"],
    ], [1.2, 1.8, 3.6])
    h2(doc, "2.1 数据流")
    bullets(doc, [
        "传感器采集环境数据，写入环境数据表。",
        "系统根据阈值判断是否异常，异常时生成告警。",
        "管理员或系统根据告警创建农事任务。",
        "农户在农户端完成任务并上传反馈。",
        "管理员审核后关闭任务和告警，形成完整记录。",
    ])
    h2(doc, "2.2 AI 调用流程")
    bullets(doc, [
        "用户提出问题，例如“番茄大棚现在要不要浇水”。",
        "Agent 判断需要哪些数据和工具。",
        "Tool 查询环境数据、作物阶段、历史任务和知识库。",
        "RAG 补充种植知识，Prompt 组织上下文。",
        "AI 输出建议，并记录调用日志，方便后续追踪。",
    ])
    h2(doc, "2.3 权限设计")
    table(doc, ["角色", "能看到什么", "不能看到什么"], [
        ["管理员", "全部业务模块、用户权限、配置、统计", "无"],
        ["农户", "自己的大棚、任务、告警、AI 建议、社区", "系统管理、其他农户数据、后台配置"],
        ["专家", "病虫害记录、AI 建议复核、社区答疑", "系统底层配置"],
    ], [1.1, 2.6, 2.9])

    h1(doc, "第3章 LangChain 智能中枢")
    para(doc, "LangChain 不直接替代业务系统，而是作为智能中枢，负责理解问题、组织上下文、调用工具、查询知识库和生成建议。")
    table(doc, ["模块", "要做什么"], [
        ["Agent", "判断用户意图，决定调用哪些 Tool，最后组织回答。"],
        ["Tool", "封装可调用能力，例如查环境数据、查任务、查告警、查作物档案。"],
        ["Memory", "保存对话上下文和用户偏好，例如用户正在问哪个大棚。"],
        ["RAG 知识库", "保存作物种植知识、病虫害处理、设备说明、项目文档。"],
        ["Prompt", "规定 AI 回答格式，要求给出原因、建议、注意事项。"],
        ["Workflow", "把复杂流程固定下来，例如告警分析、病虫害处理建议。"],
        ["AI 日志", "记录问题、调用工具、输入输出、耗时和错误，便于调试。"],
    ], [1.4, 5.0])

    h1(doc, "第4章 管理员端")
    para(doc, "管理员端是项目的后台，重点是“管理”和“统计”。")
    table(doc, ["模块", "具体功能"], [
        ["首页", "展示大棚数量、在线设备、未处理告警、今日任务、环境趋势和作物结构。"],
        ["种植管理", "管理大棚、作物档案、种植批次、生长阶段和长势记录。"],
        ["环境监测", "查看实时环境数据、历史曲线、采集记录，并配置阈值。"],
        ["告警中心", "查看异常告警，派发任务，填写处理结果，关闭告警。"],
        ["农事任务", "新建任务、派发任务、审核任务、查看完成率和逾期情况。"],
        ["AI 分析中心", "根据环境、作物和任务记录生成分析建议，供管理员审核。"],
        ["病虫害", "查看识别记录、人工复核、生成处理建议和防治任务。"],
        ["溯源", "维护批次溯源档案，生成二维码，配置公开展示内容。"],
        ["社区", "发布项目动态、审核帖子、专家答疑、沉淀知识内容。"],
    ], [1.25, 5.15])

    h1(doc, "第5章 农户端")
    para(doc, "农户端要比管理员端简单，重点是“看得懂、能操作、少配置”。")
    table(doc, ["模块", "具体功能"], [
        ["农户首页", "查看今日任务、我的告警、负责大棚状态和通知消息。"],
        ["我的大棚", "查看自己负责的大棚、作物批次、当前环境和长势记录。"],
        ["我的任务", "接收任务、开始处理、完成反馈、上传照片和查看历史任务。"],
        ["环境查看", "查看温度、湿度、土壤湿度、光照和简单趋势，不提供复杂配置。"],
        ["病虫害识别", "上传照片，查看识别结果，提交专家复核。"],
        ["AI 助手", "用简单语言回答浇水、施肥、通风、病虫害等问题。"],
        ["社区", "发布问题、查看专家回复、浏览种植经验。"],
        ["我的溯源记录", "查看自己参与的批次记录和采收信息。"],
    ], [1.25, 5.15])

    h1(doc, "第6章 数据库与 API")
    h2(doc, "6.1 主要数据表")
    bullets(doc, [
        "基地表、大棚表、作物档案表、种植批次表。",
        "设备表、传感器数据表、告警表、农事任务表、任务记录表。",
        "病虫害识别记录表、溯源批次表、社区帖子表、社区评论表。",
        "AI 调用日志表、知识库文档表、向量索引表。",
    ])
    h2(doc, "6.2 API 分类")
    table(doc, ["接口分类", "说明"], [
        ["用户与权限接口", "登录、角色、菜单、用户信息。"],
        ["业务接口", "大棚、作物、环境、告警、任务、溯源、社区。"],
        ["AI 接口", "AI 问答、病虫害识别、知识库检索、AI 日志。"],
        ["文件接口", "上传作物图片、任务反馈图片、溯源二维码。"],
    ], [1.5, 4.9])

    h1(doc, "第7章 部署与开发规范")
    bullets(doc, [
        "开发阶段先使用单体 Spring Boot 服务，MySQL 保存业务数据。",
        "AI 能力可以先用独立 Python 服务，后端通过 HTTP 调用。",
        "提交代码前先本地运行，确认登录、菜单、页面和核心接口没有报错。",
        "Git 分支建议：main 保存稳定版本，dev 做日常开发，每个人功能用 feature/模块名。",
        "提交信息写清楚，例如 feat: 新增农户任务页面，fix: 修复环境数据查询异常。",
    ])

    h1(doc, "第8章 五人开发分工")
    table(doc, ["成员", "负责内容", "交付物"], [
        ["成员1", "后台基础与权限", "登录、角色菜单、管理员/农户权限隔离、基础配置。"],
        ["成员2", "管理员端业务", "种植管理、环境监测、告警中心、农事任务。"],
        ["成员3", "农户端页面", "农户首页、我的大棚、我的任务、环境查看。"],
        ["成员4", "AI 与知识库", "LangChain、Tool、RAG、Prompt、AI 日志、AI 助手。"],
        ["成员5", "病虫害/溯源/社区/文档", "病虫害识别流程、溯源二维码、社区、演示文档。"],
    ], [0.8, 2.7, 3.1])

    para(doc, "建议开发顺序：先做权限和双端菜单，再做管理员端 P0 业务模块，然后做农户端任务闭环，最后接入 AI、病虫害、RAG 和社区能力。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "智慧农业项目开发说明", size=9, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
